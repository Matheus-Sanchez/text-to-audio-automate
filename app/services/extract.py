from __future__ import annotations

from pathlib import Path
from typing import Iterable

from app.core.config import AppSettings
from app.core.models import ExtractedDocument, Section
from app.core.text_utils import detect_language, parse_title


class DocumentExtractor:
    def __init__(self, settings: AppSettings) -> None:
        self.settings = settings

    def extract(self, path: Path) -> ExtractedDocument:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            return self._extract_text(path)
        if suffix == ".docx":
            return self._extract_docx(path)
        if suffix == ".epub":
            return self._extract_epub(path)
        if suffix == ".pdf":
            return self._extract_pdf(path)
        raise ValueError(f"Unsupported file type: {suffix}")

    def _extract_text(self, path: Path) -> ExtractedDocument:
        raw = path.read_text(encoding="utf-8", errors="ignore")
        title = parse_title(raw, path)
        sections = self._sections_from_blocks(raw)
        return ExtractedDocument(
            source_path=path,
            source_type=path.suffix.lower(),
            title=title,
            raw_text=raw,
            sections=sections,
            language_hint=detect_language(raw),
            pages_or_sections=len(sections),
        )

    def _extract_docx(self, path: Path) -> ExtractedDocument:
        import docx

        document = docx.Document(str(path))
        parts: list[str] = []
        sections: list[Section] = []
        current_heading: str | None = None
        current_lines: list[str] = []
        order = 0

        def flush_current() -> None:
            nonlocal current_heading, current_lines, order
            content = "\n".join(line for line in current_lines if line.strip()).strip()
            if content:
                sections.append(Section(heading=current_heading, content=content, order=order))
                parts.append((f"## {current_heading}\n\n" if current_heading else "") + content)
                order += 1
            current_lines = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            if style_name.startswith("heading"):
                flush_current()
                current_heading = text
            else:
                current_lines.append(text)
        flush_current()

        raw = "\n\n".join(parts).strip()
        return ExtractedDocument(
            source_path=path,
            source_type=path.suffix.lower(),
            title=parse_title(raw, path),
            raw_text=raw,
            sections=sections or self._sections_from_blocks(raw),
            language_hint=detect_language(raw),
            pages_or_sections=len(sections) or None,
        )

    def _extract_epub(self, path: Path) -> ExtractedDocument:
        from bs4 import BeautifulSoup
        from ebooklib import ITEM_DOCUMENT, epub

        book = epub.read_epub(str(path))
        parts: list[str] = []
        sections: list[Section] = []
        order = 0

        for item in book.get_items_of_type(ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_body_content(), "html.parser")
            for tag in soup(["script", "style"]):
                tag.decompose()
            heading = soup.find(["h1", "h2", "h3"])
            title = heading.get_text(" ", strip=True) if heading else None
            text = soup.get_text("\n", strip=True)
            if not text:
                continue
            sections.append(Section(heading=title, content=text, order=order, source_ref=item.file_name))
            parts.append((f"## {title}\n\n" if title else "") + text)
            order += 1

        raw = "\n\n".join(parts).strip()
        return ExtractedDocument(
            source_path=path,
            source_type=path.suffix.lower(),
            title=parse_title(raw, path),
            raw_text=raw,
            sections=sections or self._sections_from_blocks(raw),
            language_hint=detect_language(raw),
            pages_or_sections=len(sections) or None,
        )

    def _extract_pdf(self, path: Path) -> ExtractedDocument:
        import fitz
        from PIL import Image
        import pymupdf4llm
        import pytesseract

        warnings: list[str] = []
        raw_parts: list[str] = []
        sections: list[Section] = []
        ocr_sections: list[Section] = []

        try:
            markdown_text = pymupdf4llm.to_markdown(str(path))
            if markdown_text.strip():
                raw_parts.append(markdown_text.strip())
        except Exception as exc:
            warnings.append(f"pymupdf4llm_failed:{exc}")

        pytesseract.pytesseract.tesseract_cmd = str(self.settings.tesseract.executavel)
        document = fitz.open(str(path))
        for page_index, page in enumerate(document, start=1):
            text = page.get_text("text").strip()
            if len(text) >= self.settings.tesseract.fallback_chars_minimos:
                sections.append(Section(heading=f"Pagina {page_index}", content=text, order=page_index - 1, source_ref=str(page_index)))
                continue

            ocr_text = ""
            try:
                pixmap = page.get_pixmap(dpi=220, alpha=False)
                image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
                ocr_text = pytesseract.image_to_string(image, lang=self.settings.tesseract.idioma).strip()
            except Exception as exc:
                warnings.append(f"ocr_failed_page_{page_index}:{exc}")

            if ocr_text:
                warnings.append(f"ocr_used_page_{page_index}")
                section = Section(heading=f"Pagina {page_index}", content=ocr_text, order=page_index - 1, source_ref=str(page_index))
                sections.append(section)
                ocr_sections.append(section)
            elif text:
                sections.append(Section(heading=f"Pagina {page_index}", content=text, order=page_index - 1, source_ref=str(page_index)))

        if not raw_parts:
            raw_parts.extend(self._render_sections(sections))
        elif ocr_sections:
            raw_parts.append("\n\n".join(self._render_sections(ocr_sections)))

        raw = "\n\n".join(part for part in raw_parts if part.strip()).strip()
        if not sections:
            sections = self._sections_from_blocks(raw)
        return ExtractedDocument(
            source_path=path,
            source_type=path.suffix.lower(),
            title=parse_title(raw, path),
            raw_text=raw,
            sections=sections,
            language_hint=detect_language(raw),
            pages_or_sections=document.page_count,
            warnings=warnings,
            metadata={"page_count": document.page_count},
        )

    @staticmethod
    def _sections_from_blocks(text: str) -> list[Section]:
        blocks = [block.strip() for block in text.split("\n\n") if block.strip()]
        return [Section(heading=None, content=block, order=index) for index, block in enumerate(blocks)]

    @staticmethod
    def _render_sections(sections: Iterable[Section]) -> list[str]:
        rendered: list[str] = []
        for section in sections:
            prefix = f"## {section.heading}\n\n" if section.heading else ""
            rendered.append(prefix + section.content)
        return rendered
